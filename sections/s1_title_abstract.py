from utils.formatting import add_main_heading, add_formatted_body
from utils.ai_engine import get_groq_content
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def run(doc, data, client):
    # 1. TITLE PAGE (No initial page break)
    p1 = doc.add_paragraph("\n\n\n")
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p1.add_run(data['project_title'].upper())
    run1.font.size, run1.bold = Pt(24), True
    
    doc.add_paragraph("\n\n\n\n\n")
    p2 = doc.add_paragraph(f"SUBMITTED BY\n\n{data['student_name'].upper()}")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.runs[0].font.size, p2.runs[0].bold = Pt(18), True

    # 2. ABSTRACT (Strict 1 Page)
    add_main_heading(doc, "ABSTRACT", force_break=True)
    content = get_groq_content(client, f"Professional Abstract for {data['project_title']}", 280)
    add_formatted_body(doc, content)