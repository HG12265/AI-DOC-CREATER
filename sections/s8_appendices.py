from utils.formatting import add_main_heading, add_sub_heading, add_formatted_body
from utils.gemini_engine import get_gemini_diagram
from utils.ai_engine import get_groq_content # Explanation-ku Groq-e pothum
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def run(doc, data, client):
    # --- 9. APPENDICES ---
    add_main_heading(doc, "9. APPENDICES", force_break=True)
    
    diagram_list = [
        ("A. DATA FLOW DIAGRAM", "DATA FLOW DIAGRAM"),
        ("B. USE CASE DIAGRAM", "USE CASE DIAGRAM"),
        ("C. SEQUENCE DIAGRAM", "SEQUENCE DIAGRAM")
    ]
    
    for i, (title, dtype) in enumerate(diagram_list):
        if i > 0:
            doc.add_page_break()
            
        add_sub_heading(doc, title)
        
        # CALL GEMINI ENGINE
        img_stream = get_gemini_diagram(data['project_title'], dtype, data['code_ctx'])
        
        if img_stream:
            try:
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_img = p_img.add_run()
                run_img.add_picture(img_stream, width=Inches(5.5))
            except:
                doc.add_paragraph(f"[Image Render Error: {title}]")
        else:
            # Table-based fallback (Ithu eppovum break aagathu)
            doc.add_paragraph(f"\n[ DIAGRAM PLACEHOLDER: {title} ]\n")

        # Technical Explanation using Groq (Faster)
        expl_prompt = f"Professional 10-line technical explanation for '{title}' of project '{data['project_title']}'. No markdown."
        explanation = get_groq_content(client, expl_prompt, 120)
        
        doc.add_paragraph("\n")
        add_formatted_body(doc, explanation)