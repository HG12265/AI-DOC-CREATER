import re
from docx.shared import Pt, RGBColor, Inches, Mm, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def setup_page_layout(doc):
    """Sets A4 size and Double Line Borders"""
    section = doc.sections[0]
    section.page_height, section.page_width = Mm(297), Mm(210)
    
    sect_pr = section._sectPr
    pg_borders = OxmlElement('w:pgBorders')
    pg_borders.set(qn('w:offsetFrom'), 'page')
    for border in ['top', 'left', 'bottom', 'right']:
        elm = OxmlElement(f'w:{border}')
        elm.set(qn('w:val'), 'double'); elm.set(qn('w:sz'), '12'); elm.set(qn('w:space'), '24')
        pg_borders.append(elm)
    sect_pr.append(pg_borders)

def add_main_heading(doc, text, force_break=True):
    """Main Heading: 16pt, Centered, Bold, New Page"""
    if force_break: doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text.upper())
    run.font.name, run.font.size, run.bold = 'Times New Roman', Pt(16), True
    run.font.color.rgb = RGBColor(0, 51, 102)

def add_sub_heading(doc, text):
    """Sub Heading: 14pt, Left, Bold"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name, run.font.size, run.bold = 'Times New Roman', Pt(14), True

def add_formatted_body(doc, text):
    """Paragraph: 12pt, Justified, 0.4 indent, 1.5 spacing"""
    paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
    for line in paragraphs:
        is_bullet = re.match(r'^(\*|-|\d+\.)\s', line)
        if is_bullet:
            clean_line = re.sub(r'^(\*|-|\d+\.)\s', '', line)
            p = doc.add_paragraph(clean_line, style='List Bullet')
        else:
            p = doc.add_paragraph()
            if ':' in line[:30]: # Module Title Bold Logic
                title, content = line.split(':', 1)
                t_run = p.add_run(title + ":"); t_run.bold = True
                p.add_run(content)
            else: p.add_run(line)
            p.paragraph_format.first_line_indent = Inches(0.4)
            
        p.alignment, p.paragraph_format.line_spacing = WD_ALIGN_PARAGRAPH.JUSTIFY, 1.5
        for run in p.runs:
            run.font.name, run.font.size = 'Times New Roman', Pt(12)

def create_pro_table(doc, data, col_widths):
    """Creates a neat table with fixed height and centered text"""
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style, table.alignment = 'Table Grid', WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(data):
        row = table.rows[i]
        row.height = Cm(0.8)
        for j, text in enumerate(row_data):
            cell = row.cells[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = text
            cell.width = Inches(col_widths[j])
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.name, run.font.size = 'Times New Roman', Pt(11)
            if i == 0: run.bold = True